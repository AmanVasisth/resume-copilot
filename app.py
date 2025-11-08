import streamlit as st
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from huggingface_hub import InferenceClient
import traceback

DEBUG = True  # Set False for production

# ----------------------------
# 🧠 Hugging Face Model Setup with Auto Fallback
# ----------------------------
try:
    client = InferenceClient("tiiuae/falcon-7b-instruct", token=st.secrets["HUGGINGFACE_API_TOKEN"])
    # quick test ping
    client.text_generation("Hello", max_new_tokens=1)
except Exception:
    st.warning("⚠️ Falcon-7B model not available — switching to Flan-T5 (faster free model).")
    client = InferenceClient("google/flan-t5-large", token=st.secrets["HUGGINGFACE_API_TOKEN"])

# ----------------------------
# 🎨 Streamlit Page Setup
# ----------------------------
st.set_page_config(page_title="ResumeCopilot 🇮🇳", page_icon="📄", layout="wide")

st.title("📄 ResumeCopilot – AI Resume & Cover Letter Builder 🇮🇳")
st.write("Create or update your **ATS-friendly resume & cover letter** instantly using free AI from Hugging Face.")
st.markdown("---")

# ----------------------------
# 🧾 Candidate Information
# ----------------------------
st.subheader("🧾 Candidate Details")

col1, col2 = st.columns(2)
with col1:
    full_name = st.text_input("👤 Full Name")
    email = st.text_input("📧 Email Address")
    phone = st.text_input("📞 Phone Number")

with col2:
    linkedin = st.text_input("🔗 LinkedIn Profile (optional)")
    job_title = st.text_input("💼 Job Title")
    experience = st.text_input("⏳ Years of Experience")

skills = st.text_area("🧠 Key Skills (comma-separated)")
education = st.text_area("🎓 Education Background")
achievements = st.text_area("🏆 Achievements (optional)")
job_description = st.text_area("📋 Job Description (optional)", placeholder="Paste a job description here to tailor your resume...")

st.markdown("---")

# ----------------------------
# 📤 Upload Old Resume (optional)
# ----------------------------
st.subheader("📤 Upload Your Old Resume")
uploaded_file = st.file_uploader("Upload your past resume (PDF or DOCX)", type=["pdf", "docx"])
past_resume_text = ""

if uploaded_file is not None:
    if uploaded_file.type == "application/pdf":
        pdf_reader = PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            past_resume_text += page.extract_text() + "\n"
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            past_resume_text += para.text + "\n"
    st.success("✅ Resume uploaded successfully!")

st.markdown("---")

# ----------------------------
# 🚀 Generate Resume & Cover Letter
# ----------------------------
if st.button("🚀 Generate New Resume & Cover Letter"):
    with st.spinner("Creating your new AI-optimized resume... please wait ⏳"):

        prompt = f"""
        You are ResumeCopilot, an AI that creates modern, ATS-friendly resumes and cover letters.

        Candidate Details:
        - Full Name: {full_name}
        - Email: {email}
        - Phone: {phone}
        - LinkedIn: {linkedin}
        - Job Title: {job_title}
        - Experience: {experience} years
        - Skills: {skills}
        - Education: {education}
        - Achievements: {achievements}

        Job Description (optional): {job_description}

        Past Resume (if any): {past_resume_text}

        Task:
        1. Write a strong, ATS-friendly resume tailored to the job description.
        2. Include clear sections: Summary, Experience, Education, Skills, Achievements, and a Cover Letter.
        3. Use professional tone, bullet points, and concise language.
        """

        try:
            # ----------------------------
            # 🧠 Generate AI Output
            # ----------------------------
            try:
                response = client.text_generation(
                    prompt,
                    max_new_tokens=800,
                    temperature=0.7,
                    top_p=0.9
                )
                result = response
            except Exception:
                st.warning("⚠️ Primary model failed. Retrying with Flan-T5-Large (lightweight fallback).")
                fallback_client = InferenceClient("google/flan-t5-large", token=st.secrets["HUGGINGFACE_API_TOKEN"])
                response = fallback_client.text_generation(
                    prompt,
                    max_new_tokens=700,
                    temperature=0.7,
                    top_p=0.9
                )
                result = response

            # ----------------------------
            # ✅ Display Output
            # ----------------------------
            st.success("✅ Resume & Cover Letter Generated Successfully!")

            with st.expander("📋 Preview Resume & Cover Letter", expanded=True):
                st.markdown(
                    f"""
                    <div style='background-color:#f9f9f9; padding:20px; border-radius:10px; line-height:1.6;'>
                    <pre style='white-space:pre-wrap; font-family:Roboto, sans-serif; color:#333;'>{result}</pre>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            # ----------------------------
            # 💾 Create Downloadable Files
            # ----------------------------
            text_bytes = result.encode('utf-8')

            # Word File
            docx_buffer = BytesIO()
            doc = Document()
            doc.add_heading("Resume & Cover Letter", 0)
            for line in result.split("\n"):
                doc.add_paragraph(line)
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            # PDF File
            pdf_buffer = BytesIO()
            pdf = canvas.Canvas(pdf_buffer, pagesize=letter)
            width, height = letter
            y = height - 50
            for line in result.split("\n"):
                pdf.drawString(50, y, line)
                y -= 15
                if y < 50:
                    pdf.showPage()
                    y = height - 50
            pdf.save()
            pdf_buffer.seek(0)

            # Download buttons
            st.download_button("📄 Download as Text", data=text_bytes, file_name="ResumeCopilot.txt")
            st.download_button("📝 Download as Word (.docx)", data=docx_buffer, file_name="ResumeCopilot.docx")
            st.download_button("📕 Download as PDF", data=pdf_buffer, file_name="ResumeCopilot.pdf")

        except Exception as e:
            st.error("⚠️ Something went wrong while generating the resume.")
            if DEBUG:
                st.text(traceback.format_exc())
            st.warning("Retrying with smaller free model...")

            # Retry logic
            try:
                fallback_client = InferenceClient("google/flan-t5-large", token=st.secrets["HUGGINGFACE_API_TOKEN"])
                response = fallback_client.text_generation(
                    prompt, max_new_tokens=700, temperature=0.7, top_p=0.9
                )
                result = response
                st.success("✅ Successfully recovered with Flan-T5 model.")
                st.write(result)
            except Exception as inner_e:
                st.error("❌ ResumeCopilot could not recover. Please try again later.")
                if DEBUG:
                    st.text(traceback.format_exc())

st.markdown("---")
st.caption("✨ Built with ❤️ in India | ResumeCopilot.ai (Free Hugging Face Edition)")
